from pondo3 import pondo_all
import os
import sys
import numpy as np
import pandas as pd
import matplotlib as mpl
import matplotlib.pyplot as plt

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches

mpl.style.use('ggplot')
plt.rcParams["font.sans-serif"] = ["Microsoft Yahei"]
plt.rcParams["axes.unicode_minus"] = False

print(sys.path[0])
print(os.getcwd())  # 获得当前工作目录
print(os.path.abspath('.'))  # 获得当前工作目录
print(os.path.abspath('..'))  # 获得当前工作目录的父目录
print(os.path.abspath(os.curdir))  # 获得当前工作目录


line = int(sys.argv[1])
root_dir = sys.argv[2].replace("\\", "/")
data_dir = os.path.join(root_dir, "data")
coil_id_col = "coil_id"

coil_id_table_name = "coil_id_table.xlsx"
task_table_name = sys.path[0] + "/tasks/task_table_cobble.xlsx"

setup_dict = {}
setup_dict["line"] = line
setup_dict["root_dir"] = root_dir
setup_dict["coil_id_table_name"] = coil_id_table_name
setup_dict["task_table_name"] = task_table_name
setup_dict["data_dir"] = data_dir

# build coil_id_table
coil_id_list = os.listdir(data_dir)
df = pd.DataFrame()
df[coil_id_col] = coil_id_list
df.to_excel(os.path.join(root_dir, coil_id_table_name))

# columns dump in English
setup_dict["coil_id_col"] = coil_id_col
setup_dict["date_col"] = ""

result_dir = os.path.join(root_dir, "curve_matrix")
setup_dict["result_dir"] = result_dir

print(setup_dict)
pondo_all(setup_dict)


# data_dir = "/".join([root_dir, "data"])
# array_dir = "/".join([root_dir, "array"])
pic_dir = "/".join([root_dir, "plot"])
for d in [pic_dir]:
    if not os.path.exists(d):
        os.makedirs(d)

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~


def add_one_table(document, summary):
    row_num = len(summary.index) + 1
    col_num = len(summary.columns) + 1
    print(row_num, col_num)
    table = document.add_table(rows=row_num, cols=col_num)
    # 表格第一行为DataFrame的列名
    hdr_cells = table.rows[0].cells
    side_cells = table.columns[0].cells
    i = 1
    for val in summary.columns:
        hdr_cells[i].text = str(val)
        i = i + 1

    i = 1
    for val in summary.index:
        side_cells[i].text = str(val)
        i = i + 1

    i = 1
    for col in summary.columns:
        j = 1
        for idx in summary.index:
            table.columns[i].cells[j].text = (
                str(summary.loc[idx, col])
                # 注意cell的text只能接收字符串
            )
            j = j + 1
        i = i + 1
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    document.add_paragraph("")


def scale_zoom(df):
    try:
        MAX_IDX = 300
        zoom_df = pd.DataFrame()
        fp = np.linspace(0, df.loc[df["F7"].notnull()].shape[0], MAX_IDX)
        for std in std_list:
            xp = np.linspace(
                0,
                df.loc[df["F{}".format(std)].notnull()].shape[0] - 1,
                MAX_IDX)
            print(df.loc[df["F{}".format(std)].notnull()].shape)
            for idx in xp:
                z_idx = np.interp(idx, xp, fp)
                zoom_df.loc[int(z_idx), "F{}".format(
                    std)] = df.loc[int(idx), "F{}".format(std)]
    except Exception:
        zoom_df = df
    return zoom_df
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~


doc = docx.Document(sys.path[0] + "/base.docx")
doc.add_heading("堆钢和起浪调查情况", 1)

fst_std = 1
lst_std = 7
std_list = [i for i in range(fst_std, lst_std + 1)]
looper_list = [i for i in range(fst_std, lst_std)]
upstream_list = [i for i in range(fst_std, 4 + 1)]
downstream_list = [i for i in range(5, lst_std + 1)]

# ==========================================================

doc.add_heading("预摆调平情况分析", 2)

data_dict = {}
for std in std_list:
    data_dict["{}".format(std)] = pd.read_excel(
        result_dir + "/leveling{}.xlsx".format(std))

summary = pd.DataFrame()
for coil_id in coil_id_list:
    for std in std_list:
        summary.loc[coil_id, "F{}".format(std)] = round(
            data_dict["{}".format(std)][coil_id][5], 2)
doc.add_paragraph("批量的预摆调平情况如下表所示。")
add_one_table(doc, summary)

# ==========================================================
doc.add_heading("窜辊情况分析", 2)
data_dict = {}
for std in std_list:
    data_dict["{}".format(std)] = pd.read_excel(
        result_dir + "/shift_f{}.xlsx".format(std))

summary = pd.DataFrame()
for coil_id in coil_id_list:
    for std in std_list:
        summary.loc[coil_id, "F{}".format(std)] = round(
            data_dict["{}".format(std)][coil_id].mean(), 1)
doc.add_paragraph("批量的窜辊情况如下表所示。")
add_one_table(doc, summary)


# ================ single data plot =========================
single_data_list = ["fm_center_ofs", "r2_center_ofs",
                    "r2dt", "fet", "fdt", "speed7"]

for single_data in single_data_list:
    summary = pd.read_excel(
        result_dir + "/{}.xlsx".format(single_data))

    plt.figure(figsize=(10, 10))
    summary.iloc[:, -4:].plot()

    if single_data == "r2dt":
        plt.ylim(1000, 1100)

    plt.title("{}情况".format(single_data))
    pic_file = "/".join([pic_dir, "{}.png".format(single_data)])
    plt.savefig(pic_file, dpi=200)
    plt.close(0)

    doc.add_paragraph("{}情况如下图所示。".format(single_data))
    doc.add_picture(pic_file, width=Inches(5.2))


doc.save("/".join([root_dir, "调查情况.docx"]))
